import io
import logging
import pypdf
import zipfile
import xml.etree.ElementTree as ET
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)

class TextExtractor:
    """Service for extracting text from various file formats"""
    
    def __init__(self):
        self.supported_types = {
            'application/pdf': self._extract_from_pdf,
            'application/epub+zip': self._extract_from_epub,
            'text/plain': self._extract_from_txt,
            'text/markdown': self._extract_from_txt,
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document': self._extract_from_docx
        }
    
    async def extract_text(
        self, 
        file_content: bytes, 
        filename: str, 
        content_type: str
    ) -> str:
        """Extract text from file content based on content type"""
        
        logger.info(f"Extracting text from {filename} (type: {content_type})")
        
        if content_type not in self.supported_types:
            raise ValueError(f"Unsupported file type: {content_type}")
        
        try:
            extractor = self.supported_types[content_type]
            text = await extractor(file_content, filename)
            
            if not text or len(text.strip()) == 0:
                raise ValueError("No text content found in file")
            
            logger.info(f"Successfully extracted {len(text)} characters from {filename}")
            return text
            
        except Exception as e:
            logger.error(f"Text extraction failed for {filename}: {e}")
            raise Exception(f"Failed to extract text: {str(e)}")
    
    async def _extract_from_pdf(self, file_content: bytes, filename: str) -> str:
        """Extract text from PDF file"""
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = pypdf.PdfReader(pdf_file)
            
            text_content = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            if not text_content:
                raise ValueError("No readable text found in PDF")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"PDF extraction failed: {e}")
            raise Exception(f"Failed to extract PDF text: {str(e)}")
    
    async def _extract_from_epub(self, file_content: bytes, filename: str) -> str:
        """Extract text from EPUB file"""
        
        try:
            epub_file = io.BytesIO(file_content)
            
            with zipfile.ZipFile(epub_file, 'r') as zip_file:
                # Find all HTML/XHTML files in the EPUB
                html_files = [
                    name for name in zip_file.namelist() 
                    if name.endswith(('.html', '.xhtml', '.htm'))
                ]
                
                if not html_files:
                    raise ValueError("No readable content files found in EPUB")
                
                text_content = []
                
                for html_file in html_files:
                    try:
                        content = zip_file.read(html_file).decode('utf-8')
                        text = self._extract_text_from_html(content)
                        if text:
                            text_content.append(text)
                    except Exception as e:
                        logger.warning(f"Failed to extract from {html_file}: {e}")
                        continue
                
                if not text_content:
                    raise ValueError("No readable text found in EPUB")
                
                return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"EPUB extraction failed: {e}")
            raise Exception(f"Failed to extract EPUB text: {str(e)}")
    
    async def _extract_from_txt(self, file_content: bytes, filename: str) -> str:
        """Extract text from plain text file"""
        
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    text = file_content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise ValueError("Could not decode text file with any supported encoding")
            
        except Exception as e:
            logger.error(f"Text file extraction failed: {e}")
            raise Exception(f"Failed to extract text file: {str(e)}")
    
    async def _extract_from_docx(self, file_content: bytes, filename: str) -> str:
        """Extract text from DOCX file"""
        
        try:
            docx_file = io.BytesIO(file_content)
            doc = DocxDocument(docx_file)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            if not text_content:
                raise ValueError("No readable text found in DOCX")
            
            return '\n\n'.join(text_content)
            
        except Exception as e:
            logger.error(f"DOCX extraction failed: {e}")
            raise Exception(f"Failed to extract DOCX text: {str(e)}")
    
    def _extract_text_from_html(self, html_content: str) -> str:
        """Extract plain text from HTML content"""
        
        try:
            # Simple HTML text extraction using regex
            import re
            
            # Remove script and style elements
            html_content = re.sub(r'<script.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            html_content = re.sub(r'<style.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
            
            # Remove HTML tags
            text = re.sub(r'<[^>]+>', '', html_content)
            
            # Decode HTML entities
            import html
            text = html.unescape(text)
            
            # Clean up whitespace
            text = re.sub(r'\s+', ' ', text)
            text = text.strip()
            
            return text
            
        except Exception as e:
            logger.warning(f"HTML text extraction failed: {e}")
            return ""
    
    def get_supported_types(self) -> list:
        """Get list of supported file types"""
        return list(self.supported_types.keys())
    
    def is_supported_type(self, content_type: str) -> bool:
        """Check if content type is supported"""
        return content_type in self.supported_types